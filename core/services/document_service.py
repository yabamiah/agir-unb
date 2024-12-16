from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT

from pathlib import Path

import pypandoc

class DocumentService:
    """
    A service to generate personalized Word documents.
    """

    def __init__(self, output_dir: str = "data/documents"):
        """
        Initializes the service with an optional output directory.

        Args:
            output_dir (str): Directory where documents will be saved. Defaults to 'data/documents'.
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def create_document(
        self,
        title: str,
        content: list[str],
        footer: str = None,
        file_name: str = "document.docx",
    ) -> str:
        """
        Creates a personalized Word document.

        Args:
            title (str): Title of the document.
            content (list[str]): A list of paragraphs to include in the document.
            footer (str): Optional footer text.
            file_name (str): Name of the output Word file.

        Returns:
            str: Path to the generated document.
        """
        document = Document()

        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for paragraph in content:
            document.add_paragraph(paragraph)

        if footer:
            footer_section = document.sections[-1]
            footer_paragraph = footer_section.footer.paragraphs[0]
            footer_paragraph.text = footer
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        output_path = self.output_dir / file_name
        document.save(output_path)

        return str(output_path)
    
    def add_table(self, doc: Document, data: list[list[str]]) -> None:
        """
        Adds a table to the given Word document.

        Args:
            doc (Document): An existing Word document object.
            data (list[list[str]]): Data for the table. Each inner list represents a row.
        """
        
        if not data:
            raise ValueError("No data provided for the table.")
        
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        
        for i, row in enumerate(data):
            for j, cell_data in enumerate(row):
                table.cell(i, j).text = cell_data

    def create_document_with_table(self, title: str, data: list[list[str]], filename: str) -> str:
        """
        Generates a Word document with a table.

        Args:
            title (str): The title of the document.
            data (list[list[str]]): Data for the table. Each inner list represents a row.
            filename (str): The name of the Word file to be saved (without extension).

        Returns:
            str: The path to the saved document.
        """
        
        doc = Document()

        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.add_table(doc, data)

        file_path = self.output_dir / f"{filename}.docx"
        doc.save(file_path)
        print(f"Document with table saved to {file_path}")
        return str(file_path)
    
    def generate_report(self, data: dict, file_name: str = "report.docx") -> str:
        """
        Generates a detailed report from a data dictionary.

        Args:
            data (dict): A dictionary where keys are section titles and values are content lists.
            file_name (str): Name of the output Word file.

        Returns:
            str: Path to the generated report.
        """
        
        document = Document()

        for section_title, paragraphs in data.items():
            document.add_heading(section_title, level=1)

            for paragraph in paragraphs:
                document.add_paragraph(paragraph)

        output_path = self.output_dir / file_name
        document.save(output_path)

        return str(output_path)
    
    def generate_default_report(self, content: list[str], file_name: str = "default_report.docx") -> str:
        """
        Generates a default report with a fixed title and footer.

        Args:
            content (list[str]): A list of paragraphs to include in the report.
            file_name (str): Name of the output Word file. Defaults to 'default_report.docx'.

        Returns:
            str: Path to the generated report.
        """
        
        title = "Default Report Title"
        footer = "This is a system-generated report."

        return self.create_document(
            title=title, content=content, footer=footer, file_name=file_name
        )
        
    def convert_docx_to_pdf(self, input_file: str, output_file: str = "document.docx") -> None:
        """
        Converts a .docx file to a .pdf file using Pandoc.

        Args:
            input_file (str): Path to the .docx file.
            output_file (str): Path to save the output .pdf file.
        """
        
        try:
            pypandoc.convert_file(input_file, 'pdf', outputfile= self.output_dir / output_file)
            print(f"File converted successfully to {output_file}")
        except Exception as e:
            print(f"An error occurred: {e}")
