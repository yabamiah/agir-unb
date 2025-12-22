###################################################################
## DANI - Desenvolvedor e Apresentador de Números e Indicadores ##
###################################################################
"""
DANI - Desenvolvedor e Apresentador de Números e Indicadores
Módulo principal para análise de documentos e geração de índices.
"""

import os
from os import listdir
from os.path import isfile, isdir, join
import re
import sys
import io
import json
import gc
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock, Semaphore
import hashlib
from collections import defaultdict

import time
import shutil
from docx import Document
from loguru import logger
from typing import List, Dict, Tuple, Set, Optional

# Módulos internos refatorados
from core.workers.adaptive_manager import AdaptiveWorkerManager
from core.processors.pdf_processor import OptimizedPdfProcessor, PYMUPDF_AVAILABLE, PDFPLUMBER_AVAILABLE
from core.processors.docx_converter import DocxConverter
from core.services.aws_service import S3Service

# Motor NLP para planos de integridade
from core.motor_nlp.classificador_eixos import ClassificadorEixos
from core.motor_nlp.pontuador_maturidade import PontuadorMaturidadeDANI
from core.motor_nlp.calculador_imga import CalculadorIMGA


# Removido: AdaptiveWorkerManager (movido para core/workers/adaptive_manager.py)
# Removido: OptimizedPdfProcessor (movido para core/processors/pdf_processor.py)



class Dani:
    """
    DANI - Desenvolvedor e Apresentador de Números e Indicadores
    VERSÃO OTIMIZADA com paralelismo e deduplicação eficiente
    CORREÇÕES PARA GRANDES VOLUMES DE ARQUIVOS
    """

    def __init__(self,
                 docs_path: str = None,
                 output_path: str = None,
                 docx_path: str = None,
                 max_workers: int = None,
                 result_pdf: str = None,
                 batch_size: int = 50,
                 all_orgaos: bool = None,
                 orgao_name: str = None,
                 read_ratio: int = None,
                 only_integrity_plans: bool = False,
                 keywords_file: str = None):

        base_dir = os.path.normpath(os.path.join(os.path.dirname(__file__), '..', '..'))

        self.s3_service = S3Service(logger=logger)
        
        if only_integrity_plans:
            self.docs_path = os.path.join(base_dir, 'data', 'dani', 'docs', 'integridade')
        else:
            self.docs_path = docs_path or os.path.join(base_dir, 'data', 'dani', 'docs', 'input')

        if not os.path.exists(self.docs_path):
            os.makedirs(self.docs_path)

        if not only_integrity_plans and len(os.listdir(self.docs_path)) == 0:
            self._download_from_s3()

        self.output_path = output_path or os.path.join(base_dir, 'data', 'dani', 'docs', 'output', 'resultados.txt')
        self.docx_path = docx_path or os.path.join(base_dir, 'data', 'dani', 'docs', 'output')
        self.result_pdf = result_pdf or os.path.join(base_dir, 'data', 'dani', 'docs', 'result_pdf')

        os.makedirs(os.path.dirname(self.output_path), exist_ok=True)
        os.makedirs(self.docx_path, exist_ok=True)

        # Gerenciador de workers adaptativo
        self.worker_manager = AdaptiveWorkerManager(logger)
        
        # Configurações de paralelismo OTIMIZADAS (compatibilidade)
        self.max_workers = max_workers or self.worker_manager.get_optimal_workers('pdf_extraction', 10)
        self.pdf_workers = self.worker_manager.get_optimal_workers('pdf_extraction', 5)
        self.processing_workers = self.worker_manager.get_optimal_workers('text_processing', 8)
        self.batch_size = batch_size

        # Controle de recursos adaptativo
        system_cores = self.worker_manager.system_info['cpu_logical']
        self.memory_semaphore = Semaphore(min(20, system_cores * 2))
        self.file_semaphore = Semaphore(min(100, system_cores * 10))

        self.pdf_conversion_lock = Lock()

        # Variáveis de estado (podem ser definidas via parâmetros ou catch_param)
        self.orgao_name = orgao_name or ""
        self.all_orgaos = all_orgaos if all_orgaos is not None else False
        self.read_ratio = read_ratio if read_ratio is not None else 0  # 0 = todos os documentos
        self.params_provided = all_orgaos is not None or orgao_name is not None or read_ratio is not None
        self.total_words_read = 0
        self.keywords: Dict[str, int] = {}
        self.saved_snippets: List[str] = []
        self.name_docs_read: List[str] = []
        self.docx_files = {}
        self.file_atas_dict = {}
        self.orgaos_processados = []
        self.logger = logger

        # Otimizações para deduplicação com LIMPEZA DE MEMÓRIA
        self.snippet_hashes: Set[str] = set()
        self.snippet_lock = Lock()
        self.vectorizer = None
        self.is_vectorizer_fitted = False

        # Cache com limite de tamanho
        self.similarity_cache: Dict[str, float] = {}
        self.max_cache_size = 1000  # Limita o cache

        # Compilação de regex para otimização
        self.compiled_patterns: Dict[str, re.Pattern] = {}
        
        # Processador de PDF otimizado
        self.pdf_processor = OptimizedPdfProcessor(self.logger)
        
        # Conversor DOCX para PDF
        self.docx_converter = DocxConverter(self.logger, self.pdf_workers)

        # NLP Engine Integration
        if only_integrity_plans:
            self.classificador = ClassificadorEixos()
            dicionario_path = os.path.join(base_dir, 'core', 'motor_nlp', 'dicionario.json')
            try:
                with open(dicionario_path, 'r', encoding='utf-8') as f:
                    dicionario = json.load(f)
                    self.classificador.carregar_dicionario(dicionario)
            except Exception as e:
                self.logger.error(f"Erro ao carregar dicionario.json: {e}")
            
            self.pontuador = PontuadorMaturidadeDANI()
            self.calculador = CalculadorIMGA()
            self.only_integrity_plans = True
            self.imga_results = {}
        else:
            self.only_integrity_plans = False
            self.classificador = None
            self.pontuador = None
            self.calculador = None
            self.imga_results = None
        
        # Caminho para arquivo de keywords (modo não-interativo)
        self.keywords_file = keywords_file

        self._setup_logger()
        self._setup_logger()
        self.docx_converter.check_dependencies()
        self._log_worker_status()

    def _check_conversion_dependencies(self) -> None:
        """Verifica e reporta o status das dependências para conversão DOCX→PDF"""
        self.logger.info("🔍 Verificando dependências para conversão DOCX→PDF...")
        
        libreoffice_available = self._check_libreoffice_available()
        pandoc_available = self._check_pandoc_available()
        
        if libreoffice_available:
            self.logger.info("✅ LibreOffice disponível - será usado como método principal")
        else:
            self.logger.warning("⚠️ LibreOffice não encontrado")
            self.logger.info("💡 Para instalar LibreOffice:")
            self.logger.info("   Ubuntu/Debian: sudo apt-get install libreoffice")
            self.logger.info("   CentOS/RHEL: sudo yum install libreoffice")
            self.logger.info("   macOS: brew install --cask libreoffice")
        
        if pandoc_available:
            self.logger.info("✅ Pandoc disponível - será usado como método alternativo")
        else:
            self.logger.warning("⚠️ Pandoc não encontrado")
            self.logger.info("💡 Para instalar Pandoc:")
            self.logger.info("   Ubuntu/Debian: sudo apt-get install pandoc")
            self.logger.info("   CentOS/RHEL: sudo yum install pandoc")
            self.logger.info("   macOS: brew install pandoc")
            self.logger.info("   Windows: choco install pandoc")
        
        if not libreoffice_available and not pandoc_available:
            self.logger.error("❌ Nenhuma ferramenta de conversão disponível!")
            self.logger.error("💡 Instale pelo menos uma das opções acima para converter DOCX→PDF")
        else:
            self.logger.info("🎉 Pelo menos uma ferramenta de conversão está disponível")

    def _log_optimization_status(self) -> None:
        """Loga o status das otimizações disponíveis"""
        self.logger.info("🚀 Status das Otimizações de PDF:")
        
        if PYMUPDF_AVAILABLE:
            self.logger.info("✅ PyMuPDF disponível - Extração ultra-rápida para PDFs com texto")
        else:
            self.logger.warning("⚠️ PyMuPDF não disponível - Instale com: pip install pymupdf")
        
        if PDFPLUMBER_AVAILABLE:
            self.logger.info("✅ pdfplumber disponível - Boa para PDFs estruturados")
        else:
            self.logger.warning("⚠️ pdfplumber não disponível - Instale com: pip install pdfplumber")

    def _log_worker_status(self) -> None:
        """Loga o status dos workers adaptativos"""
        self.logger.info("⚡ Status dos Workers Adaptativos:")
        
        system_info = self.worker_manager.system_info
        self.logger.info(f"  🖥️  CPU: {system_info['cpu_physical']} cores físicos, {system_info['cpu_logical']} cores lógicos")
        self.logger.info(f"  💾 Memória: {system_info['memory_gb']:.1f} GB")
        self.logger.info(f"  📊 Carga atual: CPU {system_info['cpu_percent']:.1f}%, Memória {system_info['memory_percent']:.1f}%")
        
        # Mostrar workers otimizados para diferentes tarefas
        pdf_workers = self.worker_manager.get_optimal_workers('pdf_extraction', 10)
        text_workers = self.worker_manager.get_optimal_workers('text_processing', 10)
        io_workers = self.worker_manager.get_optimal_workers('file_io', 10)
        
        self.logger.info(f"  🔧 Workers otimizados:")
        self.logger.info(f"    📄 PDF Extraction: {pdf_workers} workers")
        self.logger.info(f"    📝 Text Processing: {text_workers} workers")
        self.logger.info(f"    💾 File I/O: {io_workers} workers")
        
        # Mostrar métricas de performance se disponíveis
        if self.worker_manager.performance_metrics:
            self.logger.info("  📈 Métricas de Performance:")
            for task_type, metrics in self.worker_manager.performance_metrics.items():
                if metrics:
                    avg_duration = sum(m['duration'] for m in metrics[-10:]) / min(10, len(metrics))
                    self.logger.info(f"    {task_type}: {avg_duration:.2f}s (últimas 10 execuções)")



    def docx_to_pdf(self, docx_file_path: str, pdf_output_path: str) -> bool:
        """
        Conversão robusta de DOCX para PDF
        Delega para o DocxConverter
        """
        return self.docx_converter.docx_to_pdf(docx_file_path, pdf_output_path)

    def convert_single_docx_to_pdf(self, docx_file_path: str, pdf_output_dir: str) -> bool:
        """
        Converte um único arquivo DOCX para PDF usando a biblioteca disponível

        Args:
            docx_file_path: Caminho completo para o arquivo DOCX
            pdf_output_dir: Diretório onde o PDF será salvo

        Returns:
            bool: True se a conversão foi bem-sucedida, False caso contrário
        """
        try:
            if not os.path.exists(docx_file_path):
                self.logger.warning(f"Arquivo DOCX não encontrado: {docx_file_path}")
                return False

            # Verificar se o arquivo não está corrompido
            try:
                file_size = os.path.getsize(docx_file_path)
                if file_size == 0:
                    self.logger.warning(f"Arquivo DOCX está vazio: {docx_file_path}")
                    return False
                elif file_size < 1024:  # Menos de 1KB pode indicar arquivo corrompido
                    self.logger.warning(f"Arquivo DOCX muito pequeno (pode estar corrompido): {docx_file_path} ({file_size} bytes)")
            except OSError as e:
                self.logger.error(f"Erro ao verificar arquivo DOCX: {docx_file_path} - {str(e)}")
                return False

            os.makedirs(pdf_output_dir, exist_ok=True)

            pdf_filename = os.path.splitext(os.path.basename(docx_file_path))[0] + '.pdf'
            pdf_output_path = os.path.join(pdf_output_dir, pdf_filename)

            # Verificar se o PDF já existe
            if os.path.exists(pdf_output_path):
                self.logger.info(f"📄 PDF já existe, pulando conversão: {pdf_filename}")
                return True

            self.logger.debug(f"🔄 Convertendo: {os.path.basename(docx_file_path)} → {pdf_filename}")
            
            success = self.docx_to_pdf(docx_file_path, pdf_output_path)
            
            if success:
                # Verificar se o PDF foi realmente criado e não está vazio
                if os.path.exists(pdf_output_path) and os.path.getsize(pdf_output_path) > 0:
                    self.logger.info(f"✅ Conversão bem-sucedida: {pdf_filename}")
                    return True
                else:
                    self.logger.error(f"❌ PDF criado mas está vazio: {pdf_output_path}")
                    return False
            else:
                self.logger.error(f"❌ Falha na conversão: {os.path.basename(docx_file_path)}")
                return False
                
        except Exception as e:
            self.logger.error(f"❌ Erro inesperado ao converter {docx_file_path}: {str(e)}")
            self.logger.debug(f"Traceback: {traceback.format_exc()}")
            return False

    def convert_orgao_docx_to_pdf(self, orgao_name: str) -> Tuple[int, int, int]:
        """
        Converte todos os arquivos DOCX de um órgão específico para PDF

        Args:
            orgao_name: Nome do órgão

        Returns:
            tuple: (sucessos, falhas, total)
        """
        orgao_docx_dir = os.path.join(self.docx_path, orgao_name)
        orgao_pdf_dir = os.path.join(self.result_pdf, orgao_name)

        if not os.path.exists(orgao_docx_dir):
            self.logger.warning(f"Diretório não encontrado para {orgao_name}: {orgao_docx_dir}")
            return (0, 0, 0)

        # Listar arquivos DOCX e verificar se são válidos
        all_files = os.listdir(orgao_docx_dir)
        docx_files = []
        invalid_files = []
        
        for f in all_files:
            if f.endswith('.docx'):
                file_path = os.path.join(orgao_docx_dir, f)
                try:
                    # Verificar se o arquivo é válido
                    if os.path.getsize(file_path) > 0:
                        docx_files.append(f)
                    else:
                        invalid_files.append(f)
                        self.logger.warning(f"Arquivo DOCX vazio ignorado: {f}")
                except OSError:
                    invalid_files.append(f)
                    self.logger.warning(f"Erro ao acessar arquivo DOCX: {f}")

        if invalid_files:
            self.logger.warning(f"⚠️ {len(invalid_files)} arquivos DOCX inválidos ignorados para {orgao_name}")

        if not docx_files:
            self.logger.info(f"Nenhum arquivo DOCX válido encontrado para {orgao_name}")
            return (0, 0, 0)

        self.logger.info(f"📁 Convertendo {len(docx_files)} arquivos DOCX do órgão: {orgao_name}")

        sucessos = 0
        falhas = 0
        arquivos_ja_existentes = 0

        max_workers = min(self.pdf_workers, 3)

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_file = {}

            for docx_file in docx_files:
                docx_path = os.path.join(orgao_docx_dir, docx_file)
                pdf_filename = os.path.splitext(docx_file)[0] + '.pdf'
                pdf_path = os.path.join(orgao_pdf_dir, pdf_filename)
                
                # Verificar se o PDF já existe
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    arquivos_ja_existentes += 1
                    self.logger.debug(f"📄 PDF já existe: {pdf_filename}")
                    continue
                
                future = executor.submit(self.convert_single_docx_to_pdf, docx_path, orgao_pdf_dir)
                future_to_file[future] = docx_file

            for future in as_completed(future_to_file):
                docx_file = future_to_file[future]
                try:
                    if future.result():
                        sucessos += 1
                    else:
                        falhas += 1
                except Exception as e:
                    self.logger.error(f"❌ Erro no processamento de {docx_file}: {str(e)}")
                    self.logger.debug(f"Traceback: {traceback.format_exc()}")
                    falhas += 1

        total = len(docx_files)
        self.logger.info(f"📊 Órgão {orgao_name}: {sucessos}/{total} conversões bem-sucedidas")
        if arquivos_ja_existentes > 0:
            self.logger.info(f"📄 {arquivos_ja_existentes} PDFs já existiam e foram pulados")

        return (sucessos, falhas, total)

    def convert_all_docx_to_pdf(self) -> Dict[str, Dict[str, int]]:
        """
        Converte todos os arquivos DOCX de todos os órgãos para PDF

        Returns:
            dict: Estatísticas da conversão por órgão
        """
        if not os.path.exists(self.docx_path):
            self.logger.error(f"Diretório de DOCX não encontrado: {self.docx_path}")
            return {}

        orgaos_dirs = [d for d in os.listdir(self.docx_path)
                       if os.path.isdir(os.path.join(self.docx_path, d))]

        if not orgaos_dirs:
            self.logger.warning("Nenhum diretório de órgão encontrado")
            return {}

        self.logger.info(f"🚀 Iniciando conversão DOCX→PDF para {len(orgaos_dirs)} órgãos")
        start_time = time.time()

        estatisticas = {}
        total_sucessos = 0
        total_falhas = 0
        total_arquivos = 0

        for orgao in orgaos_dirs:
            sucessos, falhas, total = self.convert_orgao_docx_to_pdf(orgao)

            estatisticas[orgao] = {
                'sucessos': sucessos,
                'falhas': falhas,
                'total': total,
                'taxa_sucesso': (sucessos / total * 100) if total > 0 else 0
            }

            total_sucessos += sucessos
            total_falhas += falhas
            total_arquivos += total

        end_time = time.time()
        duracao = end_time - start_time

        self.logger.info("=" * 60)
        self.logger.info(f"📈 RESUMO DA CONVERSÃO DOCX → PDF")
        self.logger.info("=" * 60)
        self.logger.info(f"⏱️  Tempo total: {duracao:.2f} segundos")
        self.logger.info(f"📁 Órgãos processados: {len(orgaos_dirs)}")
        self.logger.info(f"📄 Total de arquivos: {total_arquivos}")
        self.logger.info(f"✅ Sucessos: {total_sucessos}")
        self.logger.info(f"❌ Falhas: {total_falhas}")
        self.logger.info(
            f"📊 Taxa de sucesso: {(total_sucessos / total_arquivos * 100):.1f}%" if total_arquivos > 0 else "N/A")
        self.logger.info("=" * 60)

        for orgao, stats in estatisticas.items():
            if stats['total'] > 0:
                self.logger.info(f"🏢 {orgao}: {stats['sucessos']}/{stats['total']} ({stats['taxa_sucesso']:.1f}%)")

        return estatisticas

    def _download_from_s3(self):
        """Download otimizado dos arquivos da S3 com controle de recursos"""
        diretorios = self.s3_service.listar_diretorios(bucket='agir-bucket', prefixo='dani-docs/')
        if not diretorios:
            logger.warning(f"Não foi encontrar os diretorios na S3 para {self.docs_path}")
            exit(0)

        with ThreadPoolExecutor(max_workers=min(self.max_workers, 3)) as executor:
            futures = []
            for diretorio in diretorios:
                future = executor.submit(
                    self.s3_service.download_object_by_directory,
                    'agir-bucket',
                    'dani-docs/',
                    os.path.join(self.docs_path, diretorio)
                )
                futures.append(future)

            for future in as_completed(futures):
                try:
                    if not future.result():
                        logger.warning(f"Não foi possível baixar os arquivos do bucket da S3")
                        exit(0)
                except Exception as e:
                    logger.error(f"Erro no download da S3: {e}")
                    exit(0)

    def _setup_logger(self):
        base_dir = os.path.normpath(os.path.join(os.path.dirname(__file__), '..', '..'))
        log_path = os.path.join(base_dir, 'logs', 'dani_logs.log')

        os.makedirs(os.path.dirname(log_path), exist_ok=True)

        self.logger.add(log_path,
                        rotation="1 MB",
                        retention="7 days",
                        level="INFO",
                        encoding="utf-8"
                        )
        self.logger.info(f"🚀 Iniciando DANI Otimizado com {self.max_workers} workers e batch_size {self.batch_size}...")

    def _carregar_pelo_terminal(self):
        """Função interna para carregar keywords digitadas no terminal."""
        print("-" * 50)
        tratar_keywords = input("Digite as palavras-chave separadas por vírgula: ").strip()
        return [k.strip() for k in tratar_keywords.split(',') if k.strip()]

    def _carregar_por_arquivo(self, caminho: str = None):
        """Função interna para carregar keywords de um arquivo .txt.
        
        Args:
            caminho: Caminho opcional para o arquivo. Se não fornecido, pergunta ao usuário.
        """
        if caminho is None:
            print("-" * 50)
            caminho_do_arquivo = input("Digite o caminho para o arquivo .txt com as palavras-chave: ").strip()
        else:
            caminho_do_arquivo = caminho

        try:
            with open(caminho_do_arquivo, 'r', encoding='utf-8') as arquivo:
                return [linha.strip() for linha in arquivo if linha.strip()]
        except FileNotFoundError:
            self.logger.error(f"🚨 Erro: O arquivo '{caminho_do_arquivo}' não foi encontrado.")
            return []
        except Exception as e:
            self.logger.error(f"🚨 Ocorreu um erro inesperado ao ler o arquivo: {e}")
            return []

    def obter_keywords(self):
        """Obtém palavras-chave. Se keywords_file foi fornecido, carrega automaticamente."""
        keywords_carregadas = []
        
        # Modo não-interativo: carregar de arquivo diretamente
        if self.keywords_file:
            keywords_carregadas = self._carregar_por_arquivo(self.keywords_file)
        else:
            # Modo interativo: perguntar ao usuário
            while True:
                print("\nComo você deseja fornecer as palavras-chave?")
                print("1 - Digitar diretamente no terminal")
                print("2 - Fornecer um caminho de arquivo (.txt)")

                escolha = input("Digite sua escolha (1 ou 2): ").strip()

                if escolha == '1':
                    keywords_carregadas = self._carregar_pelo_terminal()
                    break
                elif escolha == '2':
                    keywords_carregadas = self._carregar_por_arquivo()
                    break
                else:
                    print("🚨 Escolha inválida. Por favor, digite 1 ou 2.")

        if keywords_carregadas:
            self.keywords = {k.lower(): 0 for k in keywords_carregadas}
            for keyword in self.keywords.keys():
                self.compiled_patterns[keyword] = re.compile(
                    r'\b' + re.escape(keyword) + r'\w*\b',
                    re.IGNORECASE
                )
            if not self.keywords_file:
                print("\n✅ Palavras-chave carregadas com sucesso!")
                print(f"   Keywords: {list(self.keywords.keys())}")
            else:
                self.logger.info(f"✅ {len(self.keywords)} palavras-chave carregadas de {self.keywords_file}")
        else:
            if not self.keywords_file:
                print("\n⚠️ Nenhuma palavra-chave foi carregada.")
            else:
                self.logger.warning(f"⚠️ Nenhuma palavra-chave carregada de {self.keywords_file}")
            self.keywords = {}

    def run(self):
        """Método principal para execução da análise otimizada"""
        try:
            self.catch_param()

            start_time = time.time()
            logger.info("⏱️ Iniciando processamento paralelo...")

            if self.all_orgaos:
                self.read_all_docs_parallel_batched()
            else:
                self.read_docs_parallel_batched()

            #self.upload_docx_s3()
            self.close_all_docs()

            end_time = time.time()
            total_time = end_time - start_time
            logger.info(f"✅ Processamento concluído em {total_time:.2f} segundos!")

            self.display_results()

        except Exception as e:
            logger.error(f"Erro crítico na execução: {e}")
            logger.error(traceback.format_exc())
        finally:
            self._cleanup_memory()

    def run_with_pdf_conversion(self):
        """
        Método principal que inclui conversão para PDF usando bibliotecas Python
        """
        try:
            self.catch_param()

            start_time = time.time()
            self.logger.info("⏱️ Iniciando processamento paralelo...")

            if self.all_orgaos:
                self.read_all_docs_parallel_batched()
            else:
                self.read_docs_parallel_batched()

            #self.upload_docx_s3()

            self.logger.info("🔄 Iniciando conversão DOCX → PDF...")
            pdf_stats = self.convert_all_docx_to_pdf()

            self.close_all_docs()

            end_time = time.time()
            total_time = end_time - start_time
            self.logger.info(f"✅ Processamento completo (incluindo PDF) concluído em {total_time:.2f} segundos!")

            self.display_results()

            return pdf_stats

        except Exception as e:
            self.logger.error(f"Erro crítico na execução: {e}")
            self.logger.error(traceback.format_exc())
        finally:
            self._cleanup_memory()

    def catch_param(self) -> None:
        """Coleta parâmetros do usuário para a análise"""
        # Se os parâmetros já foram fornecidos no construtor, não precisa pedir
        if self.params_provided:
            self.logger.info("Parâmetros fornecidos via construtor, pulando entrada interativa")
            if self.all_orgaos:
                self.logger.info("Modo: Processar todos os órgãos")
            else:
                self.logger.info(f"Modo: Processar órgão específico - {self.orgao_name}")
                self.logger.info(f"Quantidade de documentos: {self.read_ratio if self.read_ratio > 0 else 'Todos'}")
        else:
            # Verifica se está em modo interativo (stdin disponível e é um TTY)
            is_interactive = sys.stdin.isatty() if hasattr(sys.stdin, 'isatty') else False
            
            if is_interactive:
                # Modo interativo - pede parâmetros ao usuário
                all_orgaos_input = input("Você deseja fazer uma pesquisa geral (s/n): ").strip().lower()
                self.all_orgaos = (all_orgaos_input == 's')

                if not self.all_orgaos:
                    self.orgao_name = input("Insira o nome do orgão: ").strip()
                    acronym = company_acronym(self.orgao_name.upper())
                    self.orgao_name = acronym if acronym else self.orgao_name.upper()
                    read_ratio_input = input("Insira a quantidade de documentos que serão lidos (0 para todos): ").strip()
                    self.read_ratio = int(read_ratio_input) if read_ratio_input else 0
                else:
                    self.read_ratio = 0  # Todos os documentos quando processar todos os órgãos
            else:
                # Modo não-interativo (background) - usa valores padrão
                self.logger.info("Modo não-interativo detectado, usando valores padrão")
                self.all_orgaos = True  # Por padrão, processa todos os órgãos
                self.read_ratio = 0  # Todos os documentos
                self.logger.info("Configuração padrão: Processar todos os órgãos, todos os documentos")

        # Configurar encoding do stdin apenas se estiver disponível
        try:
            if hasattr(sys.stdin, 'buffer'):
                sys.stdin = io.TextIOWrapper(sys.stdin.buffer, encoding='utf-8')
        except (AttributeError, OSError) as e:
            self.logger.debug(f"Não foi possível configurar encoding do stdin: {e}")
        
        self.obter_keywords()
        self.logger.info(f"Palavras-chave: {list(self.keywords.keys())}")

    def read_docs_parallel_batched(self) -> None:
        """Lê documentos de um órgão específico usando paralelismo em lotes"""
        docs_path = os.path.join(self.docs_path, self.orgao_name)
        files = [f for f in listdir(docs_path) if isfile(join(docs_path, f))]

        if self.read_ratio == 0:
            files_to_process = files
        else:
            files_to_process = files[:self.read_ratio]

        self.logger.info(f"📚 Processando {len(files_to_process)} arquivos em lotes de {self.batch_size}")

        for i in range(0, len(files_to_process), self.batch_size):
            batch = files_to_process[i:i + self.batch_size]
            self.logger.info(
                f"🔄 Processando lote {i // self.batch_size + 1}/{(len(files_to_process) - 1) // self.batch_size + 1}")

            self._process_batch(batch, docs_path, self.orgao_name)

            gc.collect()
            time.sleep(0.1)

        self.save_docs_for_current_orgao(self.orgao_name)

    def read_all_docs_parallel_batched(self) -> None:
        """Lê documentos de todos os órgãos usando paralelismo em lotes"""
        all_dir = [join(self.docs_path, d) for d in listdir(self.docs_path) if isdir(join(self.docs_path, d))]
        
        # Se estiver no modo integridade, não pular órgãos já processados
        if self.only_integrity_plans:
            orgaos_to_read = all_dir
        else:
            read_orgaos = [os.path.basename(d) for d in listdir(self.docx_path) if isdir(join(self.docx_path, d))]
            orgaos_to_read = [orgao for orgao in all_dir if os.path.basename(orgao) not in read_orgaos]

        self.logger.info(f"Órgãos a serem lidos: {[os.path.basename(o) for o in orgaos_to_read]}")

        all_files = []
        for dir_path in orgaos_to_read:
            orgao_atual = os.path.basename(dir_path)
            files_in_dir = [f for f in listdir(dir_path) if isfile(join(dir_path, f))]
            for file in files_in_dir:
                all_files.append((join(dir_path, file), orgao_atual, file))

        self.logger.info(f"📚 Processando {len(all_files)} arquivos em lotes de {self.batch_size}")

        orgao_files = defaultdict(list)
        for file_path, orgao_name, filename in all_files:
            orgao_files[orgao_name].append((file_path, orgao_name, filename))

        for orgao_name, files in orgao_files.items():
            self.logger.info(f"🏢 Processando órgão: {orgao_name} ({len(files)} arquivos)")
            self.orgaos_processados.append(orgao_name)

            for i in range(0, len(files), self.batch_size):
                batch = files[i:i + self.batch_size]
                self.logger.info(
                    f"🔄 Lote {i // self.batch_size + 1}/{(len(files) - 1) // self.batch_size + 1} do {orgao_name}")

                self._process_batch_all_orgaos(batch)

                gc.collect()
                time.sleep(0.1)

            self.save_docs_for_current_orgao(orgao_name)

    def _process_batch(self, batch: List[str], docs_path: str, orgao_name: str) -> None:
        """Processa um lote de arquivos com workers adaptativos"""
        start_time = time.time()
        
        # Usar gerenciador adaptativo para escolher o executor ideal
        executor = self.worker_manager.get_executor('pdf_extraction', len(batch))
        
        with executor as exec:
            futures = []
            for file in batch:
                file_path = join(docs_path, file)
                future = exec.submit(self._process_single_pdf_safe, file_path, orgao_name, file)
                futures.append(future)

            for future in as_completed(futures, timeout=300):
                try:
                    result = future.result(timeout=60)  # 1 minuto para resultado
                    if result:
                        self.name_docs_read.append(result['filename'])
                        with self.snippet_lock:
                            self.total_words_read += result['words_read']
                            for keyword, count in result['keyword_counts'].items():
                                self.keywords[keyword] += count
                except Exception as e:
                    self.logger.error(f"Erro no processamento do lote: {e}")
        
        # Log de performance
        duration = time.time() - start_time
        self.worker_manager.log_performance('pdf_extraction', duration, len(batch))

    def _process_batch_all_orgaos(self, batch: List[Tuple[str, str, str]]) -> None:
        """Processa um lote de arquivos para todos os órgãos com workers adaptativos"""
        start_time = time.time()
        
        # Usar gerenciador adaptativo para escolher o executor ideal
        executor = self.worker_manager.get_executor('pdf_extraction', len(batch))
        
        with executor as exec:
            futures = []
            for file_path, orgao_name, filename in batch:
                future = exec.submit(self._process_single_pdf_safe, file_path, orgao_name, filename)
                futures.append(future)

            for future in as_completed(futures, timeout=3000):
                try:
                    result = future.result(timeout=60)
                    if result:
                        self.name_docs_read.append(result['filename'])
                        with self.snippet_lock:
                            self.total_words_read += result['words_read']
                            for keyword, count in result['keyword_counts'].items():
                                self.keywords[keyword] += count
                except Exception as e:
                    self.logger.error(f"Erro no processamento: {e}")
        
        # Log de performance
        duration = time.time() - start_time
        self.worker_manager.log_performance('pdf_extraction', duration, len(batch))

    def _process_single_pdf_safe(self, file_path: str, orgao_name: str, filename: str) -> Optional[Dict]:
        """Processa um único PDF com controle de recursos e tratamento de erros"""
        try:
            with self.memory_semaphore:
                return self._process_single_pdf(file_path, orgao_name, filename)
        except Exception as e:
            self.logger.error(f"Erro ao processar PDF {file_path}: {e}")
            return None
        finally:
            gc.collect()

    def _process_single_pdf(self, file_path: str, orgao_name: str, filename: str) -> Optional[Dict]:
        """Processa um único PDF usando o processador otimizado"""
        try:
            # Usar o processador otimizado
            pdf_pages_text, words_read = self.pdf_processor.extract_text_fast(file_path)
            
            if not pdf_pages_text:
                self.logger.warning(f"Nenhum texto extraído de {filename}")
                return None

            keyword_counts = {k: 0 for k in self.keywords.keys()}
            processed_snippets = []

            # Processar páginas com workers adaptativos
            if len(pdf_pages_text) > 1:
                # Usar gerenciador adaptativo para processamento de texto
                executor = self.worker_manager.get_executor('text_processing', len(pdf_pages_text))
                
                with executor as exec:
                    futures = []
                    for i, page_text in enumerate(pdf_pages_text):
                        if page_text:
                            future = exec.submit(self._process_page_text, page_text, filename, orgao_name)
                            futures.append(future)
                    
                    for future in as_completed(futures):
                        try:
                            page_keyword_counts, snippets = future.result(timeout=30)
                            
                            for keyword, count in page_keyword_counts.items():
                                keyword_counts[keyword] += count
                            
                            processed_snippets.extend(snippets)
                        except Exception as e:
                            self.logger.warning(f"Erro ao processar página: {e}")
            else:
                # Para PDFs com uma página, processar sequencialmente
                for page_text in pdf_pages_text:
                    if not page_text:
                        continue

                    page_keyword_counts, snippets = self._process_page_text(page_text, filename, orgao_name)

                    for keyword, count in page_keyword_counts.items():
                        keyword_counts[keyword] += count

                    processed_snippets.extend(snippets)

            # Deduplicação otimizada
            unique_snippets = self._deduplicate_snippets_optimized(processed_snippets)

            # Escrever snippets de forma otimizada
            for snippet_info in unique_snippets:
                self._write_snippet_thread_safe(snippet_info, orgao_name)


            # Se a integração NLP estiver ativa, calcular IMGA
            imga_result = None
            if self.only_integrity_plans and self.classificador and self.pontuador and self.calculador:
                total_words_doc = words_read
                scores_por_eixo = defaultdict(list)
                evidencias_por_eixo = defaultdict(list)
                boosters_por_eixo = defaultdict(lambda: {'B1_ACAO': 0, 'B2_PERIODICIDADE': 0, 'B3_RESPONSAVEL': 0, 'B4_ARTEFATO': 0})
                segmentos_analisados = 0
                
                # Iterar sobre as páginas, segmentar e limpar
                for page_text in pdf_pages_text:
                    if not page_text: continue
                    
                    # Passo 1: Segmentação e Limpeza
                    segmentos = self._segment_and_clean(page_text)
                    
                    for segmento in segmentos:
                        segmentos_analisados += 1
                        # Classificar o segmento
                        sinais_eixos = self.classificador.classificar_texto(segmento)
                        
                        # Para cada eixo, verificar pontuação neste segmento
                        for eixo_id, sinais_config in self.classificador.dicionario_eixos.items():
                             resultado = self.pontuador.calcular_score_eixo(segmento, eixo_id, sinais_config)
                             if resultado.score_final > 0:
                                 scores_por_eixo[eixo_id].append(resultado.score_final)
                                 evidencias_por_eixo[eixo_id].extend(resultado.evidencias)
                                 # Contar boosters ativados
                                 for booster_id, ativado in resultado.detalhe_boosters.items():
                                     if ativado:
                                         boosters_por_eixo[eixo_id][booster_id] += 1

                # Consolidar evidências únicas e contagem
                estatisticas_eixos = {}
                for eixo_id in self.classificador.dicionario_eixos.keys():
                    evidencias_unicas = list(set(evidencias_por_eixo.get(eixo_id, [])))
                    estatisticas_eixos[eixo_id] = {
                        'termos_encontrados': len(evidencias_unicas),
                        'termos_lista': evidencias_unicas,  # Todos os termos encontrados
                        'total_ocorrencias': len(evidencias_por_eixo.get(eixo_id, [])),
                        'boosters': dict(boosters_por_eixo.get(eixo_id, {}))
                    }

                imga_result = self.calculador.calcular_imga_entidade(scores_por_eixo, total_words_doc)
                imga_result['estatisticas_eixos'] = estatisticas_eixos
                imga_result['metadados']['segmentos_analisados'] = segmentos_analisados
                
                # Salvar resultado do IMGA associado ao arquivo
                with self.snippet_lock:
                    if self.imga_results is not None:
                         # Extrair o órgão do nome do arquivo ou usar o atual
                         # Vamos usar o nome do arquivo para chavear por enquanto
                         self.imga_results[filename] = imga_result
                         self.logger.info(f"📊 IMGA calculado para {filename}: Global={imga_result['imga_global']}, Faixa={imga_result['faixa']}")

            return {
                'filename': filename,
                'words_read': words_read,
                'keyword_counts': keyword_counts,
                'snippets_count': len(unique_snippets),
                'imga_result': imga_result
            }

        except Exception as e:
            self.logger.error(f"Erro ao processar PDF {file_path}: {e}")
            return None
        finally:
            gc.collect()

    def _process_page_text(self, page_text: str, filename: str, orgao_name: str) -> Tuple[Dict[str, int], List[Dict]]:
        """Processa o texto de uma página com otimizações avançadas"""
        keyword_counts = {k: 0 for k in self.keywords.keys()}
        snippets = []

        # Limitar tamanho da página para evitar processamento excessivo
        if len(page_text) > 50000:
            page_text = page_text[:50000]

        # Otimização: processar todas as keywords de uma vez usando regex múltiplo
        if len(self.keywords) > 1:
            # Criar um regex combinado para todas as keywords
            combined_pattern = self._create_combined_regex()
            if combined_pattern:
                matches = list(combined_pattern.finditer(page_text))
                
                # Agrupar matches por keyword
                keyword_matches = defaultdict(list)
                for match in matches:
                    for keyword in self.keywords:
                        if keyword in match.group().lower():
                            keyword_matches[keyword].append(match)
                            break
                
                # Processar cada keyword
                for keyword, keyword_match_list in keyword_matches.items():
                    if len(keyword_match_list) > 100:
                        keyword_match_list = keyword_match_list[:100]
                    
                    keyword_counts[keyword] = len(keyword_match_list)
                    
                    for match in keyword_match_list:
                        snippet_text = self._extract_snippet_optimized(page_text, match.start(), match.end())
                        snippets.append({
                            'text': snippet_text,
                            'keyword': keyword,
                            'filename': filename,
                            'orgao': orgao_name
                        })
        else:
            # Fallback para processamento individual quando há apenas uma keyword
            for keyword in self.keywords:
                pattern = self.compiled_patterns.get(keyword)
                if not pattern:
                    continue

                matches = list(pattern.finditer(page_text))
                if len(matches) > 100:
                    matches = matches[:100]

                keyword_counts[keyword] = len(matches)

                for match in matches:
                    snippet_text = self._extract_snippet_optimized(page_text, match.start(), match.end())
                    snippets.append({
                        'text': snippet_text,
                        'keyword': keyword,
                        'filename': filename,
                        'orgao': orgao_name
                    })

        return keyword_counts, snippets

    def _create_combined_regex(self) -> Optional[re.Pattern]:
        """Cria um regex combinado para todas as keywords"""
        try:
            if not self.keywords:
                return None
            
            # Criar padrão alternativo para todas as keywords
            patterns = []
            for keyword in self.keywords:
                escaped_keyword = re.escape(keyword)
                patterns.append(f'\\b{escaped_keyword}\\w*\\b')
            
            combined_pattern = '|'.join(patterns)
            return re.compile(combined_pattern, re.IGNORECASE)
        except Exception as e:
            self.logger.warning(f"Erro ao criar regex combinado: {e}")
            return None

    def _extract_snippet_optimized(self, text: str, start_pos: int, end_pos: int, window_size: int = 300) -> str:
        """Extrai snippet otimizado ao redor de uma posição - REDUZIDO para economizar memória"""
        start_index = max(0, start_pos - window_size)
        end_index = min(len(text), end_pos + window_size)
        return text[start_index:end_index]

    def _calculate_hash(self, text: str) -> str:
        """Calcula hash rápido para deduplicação"""
        return hashlib.md5(text.encode('utf-8', errors='ignore')).hexdigest()

    def _deduplicate_snippets_optimized(self, snippets: List[Dict]) -> List[Dict]:
        """Deduplicação otimizada com controle de memória"""
        if not snippets:
            return []

        unique_snippets = []
        seen_hashes = set()

        max_snippets = min(len(snippets), 500)

        for snippet in snippets[:max_snippets]:
            snippet_text = snippet['text']
            snippet_hash = self._calculate_hash(snippet_text)

            if snippet_hash in seen_hashes:
                continue

            is_duplicate = False
            if len(unique_snippets) > 0:
                recent_snippets = unique_snippets[-5:]

                for existing_snippet in recent_snippets:
                    try:
                        similarity = self._calculate_similarity_simple(snippet_text, existing_snippet['text'])
                        if similarity >= 0.85:
                            is_duplicate = True
                            break
                    except:
                        if snippet_text.strip() == existing_snippet['text'].strip():
                            is_duplicate = True
                            break

            if not is_duplicate:
                unique_snippets.append(snippet)
                seen_hashes.add(snippet_hash)

        return unique_snippets

    def _calculate_similarity_simple(self, text1: str, text2: str) -> float:
        """Calcula similaridade simples baseada em palavras - mais eficiente"""
        try:
            words1 = set(text1.lower().split())
            words2 = set(text2.lower().split())
            intersection = words1.intersection(words2)
            union = words1.union(words2)
            return len(intersection) / len(union) if union else 0
        except:
            return 0

    def _calculate_similarity_cached(self, text1: str, text2: str) -> float:
        """Calcula similaridade com cache limitado"""
        if len(self.similarity_cache) > self.max_cache_size:
            self.similarity_cache.clear()

        cache_key = f"{self._calculate_hash(text1)[:8]}_{self._calculate_hash(text2)[:8]}"

        if cache_key in self.similarity_cache:
            return self.similarity_cache[cache_key]

        try:
            similarity = self._calculate_similarity_simple(text1, text2)
            self.similarity_cache[cache_key] = similarity
            return similarity
        except:
            return 0

    def _write_snippet_thread_safe(self, snippet_info: Dict, orgao_name: str) -> None:
        """Escreve snippet de forma thread-safe com controle de arquivos"""
        try:
            with self.file_semaphore:
                with self.snippet_lock:
                    snippet_text = snippet_info['text']
                    keyword = snippet_info['keyword']
                    filename = snippet_info['filename']

                    snippet_with_file = f'[Arquivo: {filename}]:\n\n"{snippet_text}"'

                    key = (orgao_name, keyword)
                    if key not in self.docx_files:
                        self.docx_files[key] = self.create_keyword_docx(keyword, orgao_name)

                    doc = self.docx_files[key]
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1
                    p.paragraph_format.space_after = 0
                    snippet_fixed = ''.join(c if c.isprintable() else ' ' for c in snippet_with_file)
                    p.add_run(snippet_fixed)
                    p.add_run("\n")
        except Exception as e:
            self.logger.error(f"Erro ao escrever snippet: {e}")

    def create_keyword_docx(self, keyword: str, orgao_name: str) -> Document:
        """Cria documento DOCX para uma palavra-chave específica"""
        try:
            org_dir = join(self.docx_path, orgao_name)
            os.makedirs(org_dir, exist_ok=True)
            new_docx_path = join(org_dir, f'output_{orgao_name}_{keyword}.docx')

            if isfile(new_docx_path):
                doc = Document(new_docx_path)
            else:
                doc = Document()
                doc.add_heading('Resultado Dani', 0)
                doc.add_heading(f'Análise atas CIG {orgao_name}', 2)
                doc.add_heading(f'Palavra-chave: {keyword}', 2)
            return doc
        except Exception as e:
            self.logger.error(f"Erro ao criar documento DOCX: {e}")
            return Document()

    def save_docs_for_current_orgao(self, orgao_name: str) -> None:
        """Salva documentos para o órgão atual com tratamento de erros"""
        try:
            org_dir = join(self.docx_path, orgao_name)
            os.makedirs(org_dir, exist_ok=True)
            docs_to_save = {k: v for k, v in self.docx_files.items() if k[0] == orgao_name}

            for (org, keyword), doc in docs_to_save.items():
                try:
                    file_path = join(org_dir, f'output_{org}_{keyword}.docx')
                    doc.save(file_path)
                    self.logger.info(f"Arquivo salvo: {file_path}")
                except Exception as e:
                    self.logger.error(f"Erro ao salvar arquivo {file_path}: {e}")

            keys_to_remove = [k for k in self.docx_files.keys() if k[0] == orgao_name]
            for key in keys_to_remove:
                del self.docx_files[key]

            gc.collect()

        except Exception as e:
            self.logger.error(f"Erro ao salvar documentos do órgão {orgao_name}: {e}")

    def upload_docx_s3(self):
        """Upload otimizado para S3 com paralelismo e tratamento de erros"""
        try:
            docx_path = self.docx_path
            logger.debug(f"📂 Caminho de saída: {docx_path}")

            os.makedirs(docx_path, exist_ok=True)
            if not os.path.exists(docx_path):
                logger.error("❌ Falha ao enviar docxs para o S3. DOCX não encontrados.")
                return False

            upload_tasks = []
            entries = os.listdir(docx_path)
            for entry in entries:
                logger.debug(f"📁 Diretório de órgão encontrado: {entry}")
                orgao_path = os.path.join(docx_path, entry)
                if os.path.isdir(orgao_path):
                    docxs = os.listdir(orgao_path)
                    for docx in docxs:
                        file_path = os.path.join(orgao_path, docx)
                        object_name = f'data/dani/docs/output/{entry}/{docx}'

                        try:
                            if not self.s3_service.object_exists(bucket='agir-bucket', object_name=object_name):
                                upload_tasks.append((file_path, object_name))
                            else:
                                logger.warning(f"⚠ Arquivo já existe na S3 e será ignorado: {object_name}")
                        except Exception as e:
                            logger.error(f"Erro ao verificar existência do arquivo na S3: {e}")

            if upload_tasks:
                logger.info(f"📤 Iniciando upload de {len(upload_tasks)} arquivos")
                with ThreadPoolExecutor(max_workers=min(self.max_workers, 3)) as executor:
                    futures = []
                    for file_path, object_name in upload_tasks:
                        future = executor.submit(
                            self.s3_service.upload_object,
                            'agir-bucket',
                            object_name,
                            file_path
                        )
                        futures.append(future)

                    for future in as_completed(futures, timeout=600):
                        try:
                            future.result(timeout=120)
                        except Exception as e:
                            logger.error(f"Erro no upload: {e}")

            return True
        except Exception as e:
            logger.error(f"Erro crítico no upload para S3: {e}")
            return False

    def close_all_docs(self) -> None:
        """Limpa recursos e força garbage collection"""
        try:
            self.docx_files.clear()
            self.similarity_cache.clear()
            self.snippet_hashes.clear()
            self.compiled_patterns.clear()

            gc.collect()

        except Exception as e:
            logger.error(f"Erro ao limpar recursos: {e}")

    def _cleanup_memory(self) -> None:
        """Limpeza forçada de memória"""
        try:
            if hasattr(self, 'docx_files'):
                self.docx_files.clear()
            if hasattr(self, 'similarity_cache'):
                self.similarity_cache.clear()
            if hasattr(self, 'snippet_hashes'):
                self.snippet_hashes.clear()
            if hasattr(self, 'compiled_patterns'):
                self.compiled_patterns.clear()
            if hasattr(self, 'vectorizer'):
                self.vectorizer = None

            for _ in range(3):
                gc.collect()
                time.sleep(0.1)

        except Exception as e:
            logger.error(f"Erro na limpeza de memória: {e}")

    def _calculate_orgao_statistics(self) -> Dict[str, Dict]:
        """Calcula estatísticas detalhadas por órgão"""
        orgao_stats = {}
        
        # Agrupar documentos por órgão
        docs_por_orgao = defaultdict(list)
        for doc_name in self.name_docs_read:
            # Extrair órgão do nome do documento ou usar o órgão atual
            if self.all_orgaos:
                # Para análise geral, tentar extrair órgão do nome do arquivo
                orgao = self._extract_orgao_from_filename(doc_name)
            else:
                orgao = self.orgao_name
            
            docs_por_orgao[orgao].append(doc_name)
        
        # Calcular estatísticas para cada órgão
        for orgao, docs in docs_por_orgao.items():
            # Contar palavras-chave por órgão (aproximação baseada nos documentos)
            orgao_keywords = {}
            total_orgao_words = 0
            
            # Estimar palavras lidas por órgão (distribuição proporcional)
            if self.total_words_read > 0 and len(self.name_docs_read) > 0:
                words_per_doc = self.total_words_read / len(self.name_docs_read)
                total_orgao_words = int(words_per_doc * len(docs))
            
            # Distribuir keywords proporcionalmente
            for keyword, total_count in self.keywords.items():
                if len(self.name_docs_read) > 0:
                    orgao_count = int((total_count * len(docs)) / len(self.name_docs_read))
                    orgao_keywords[keyword] = orgao_count
            
            # Calcular porcentagens por órgão
            orgao_percentages = {}
            for keyword, count in orgao_keywords.items():
                percentage = (count / total_orgao_words) * 100.0 if total_orgao_words > 0 else 0
                orgao_percentages[keyword] = round(percentage, 2)
            
            # Top keywords do órgão
            top_keywords = sorted(orgao_keywords.items(), key=lambda x: x[1], reverse=True)[:10]
            
            orgao_stats[orgao] = {
                "documentos_processados": len(docs),
                "total_palavras_lidas": total_orgao_words,
                "contagem_keywords": orgao_keywords,
                "porcentagem_keywords": orgao_percentages,
                "top_keywords": [{"keyword": k, "count": c} for k, c in top_keywords],
                "total_ocorrencias": sum(orgao_keywords.values()),
                "media_ocorrencias_por_documento": round(sum(orgao_keywords.values()) / len(docs), 2) if docs else 0
            }
        
        return orgao_stats

    def _extract_orgao_from_filename(self, filename: str) -> str:
        """Extrai o nome do órgão a partir do nome do arquivo"""
        # Mapear padrões comuns de nomes de arquivos para órgãos
        filename_lower = filename.lower()
        
        # Padrões específicos para identificar órgãos
        orgao_patterns = {
            'sema': ['sema', 'meio ambiente'],
            'jucis': ['jucis', 'junta comercial'],
            'emater': ['emater', 'assistência técnica'],
            'pcdf': ['pcdf', 'polícia civil'],
            'undf': ['undf', 'universidade'],
            'sepe': ['sepe', 'projetos especiais'],
            'ipedf': ['ipedf', 'pesquisa estatística'],
            'detran': ['detran', 'trânsito'],
            'caci': ['caci', 'casa civil'],
            'cbm': ['cbm', 'bombeiros'],
            'cgdf': ['cgdf', 'controladoria'],
            'cm': ['cm', 'casa militar'],
            'der': ['der', 'estradas'],
            'df_legal': ['df legal', 'ordem urbanística'],
            'fapdf': ['fapdf', 'apoio pesquisa'],
            'fepECS': ['fepECS', 'ensino pesquisa'],
            'fhb': ['fhb', 'hemocentro'],
            'ibram': ['ibram', 'brasília ambiental'],
            'iprevdf': ['iprevdf', 'previdência'],
            'jbb': ['jbb', 'jardim botânico'],
            'pgdf': ['pgdf', 'procuradoria'],
            'procon': ['procon', 'consumidor'],
            'seac': ['seac', 'atendimento comunidade'],
            'seagri': ['seagri', 'agricultura'],
            'seape': ['seape', 'administração penitenciária'],
            'secti': ['secti', 'ciência tecnologia'],
            'sedes': ['sedes', 'desenvolvimento social'],
            'seduh': ['seduh', 'desenvolvimento urbano'],
            'seedf': ['seedf', 'educação'],
            'segov': ['segov', 'governo'],
            'sejus': ['sejus', 'justiça cidadania'],
            'sejuv': ['sejuv', 'família juventude'],
            'sema': ['sema', 'meio ambiente'],
            'semob': ['semob', 'transporte mobilidade'],
            'sepe': ['sepe', 'projetos especiais'],
            'ses': ['ses', 'saúde'],
            'slu': ['slu', 'limpeza urbana'],
            'smdf': ['smdf', 'mulher'],
            'sodf': ['sodf', 'obras infraestrutura'],
            'sspdf': ['sspdf', 'segurança pública']
        }
        
        for orgao, patterns in orgao_patterns.items():
            for pattern in patterns:
                if pattern in filename_lower:
                    return orgao.upper()
        
        # Se não encontrar padrão específico, tentar extrair do início do nome
        parts = filename.split('_')
        if parts:
            potential_orgao = parts[0].upper()
            if len(potential_orgao) <= 10:  # Limitar tamanho para evitar nomes muito longos
                return potential_orgao
        
        return "OUTROS"

    def _segment_and_clean(self, text: str) -> List[str]:
        """
        Segmenta o texto em unidades de análise (parágrafos) e remove ruídos.
        Refina o Passo 1 do fluxo IMGA.
        """
        # 1. Limpeza de ruídos básicos
        # Remover números de página isolados (ex: " 12 ", "- 12 -")
        text = re.sub(r'\n\s*[-–]?\s*\d+\s*[-–]?\s*\n', '\n', text)
        
        # 2. Segmentação por parágrafos (quebras de linha duplas ou recuos)
        # Normalizar quebras de linha
        text = re.sub(r'\r\n', '\n', text)
        
        # Dividir por quebras duplas (parágrafos claros)
        raw_segments = re.split(r'\n\s*\n', text)
        
        cleaned_segments = []
        for seg in raw_segments:
            seg = seg.strip()
            # Ignorar segmentos muito curtos (ruído) ou irrelevantes
            if len(seg) < 20: 
                continue
            cleaned_segments.append(seg)
            
        return cleaned_segments

    def display_results(self) -> None:
        """Exibe resultados da análise com tratamento de erros e gera summary detalhado"""
        try:
            self.logger.info("\n--- Resultados DANI Otimizado ---")
            self.logger.info(f"Órgão(s) analisado(s): {', '.join(self.orgaos_processados)}")
            self.logger.info(f"Documentos lidos: {len(self.name_docs_read)}")
            self.logger.info("Quantidade de vezes que cada palavra-chave foi encontrada:")

            # Calcular estatísticas por órgão
            orgao_statistics = self._calculate_orgao_statistics()
            
            # Exibir resultados IMGA se disponíveis
            if self.only_integrity_plans and self.imga_results:
                self.logger.info("\n=== RESULTADOS IMGA (Índice de Maturidade da Governança Algorítmica) ===")
                for filename, result in self.imga_results.items():
                    self.logger.info(f"\n📄 Documento: {filename}")
                    self.logger.info(f"   🏆 IMGA Global: {result['imga_global']} ({result['faixa']})")
                    self.logger.info("   📊 Índices por Eixo:")
                    for eixo, score in result['indices_eixos'].items():
                        nome_eixo = self.classificador.obter_info_eixo(eixo) if self.classificador else eixo
                        self.logger.info(f"      - {eixo} ({nome_eixo}): {score}")
                self.logger.info("========================================================================\n")

            # Estrutura do summary melhorada
            summary = {
                "metadata": {
                    "data_analise": datetime.now().isoformat(),
                    "versao_dani": "2.0",
                    "tipo_analise": "integridade" if self.only_integrity_plans else ("geral" if self.all_orgaos else "especifica"),
                    "orgao_especifico": self.orgao_name if not self.all_orgaos else None
                },
                "resumo_geral": {
                    "orgaos_analisados": self.orgaos_processados,
                    "total_documentos_lidos": len(self.name_docs_read),
                    "total_palavras_lidas": self.total_words_read,
                    "total_keywords_encontradas": len(self.keywords),
                    "total_ocorrencias": sum(self.keywords.values()),
                    "media_ocorrencias_por_documento": round(sum(self.keywords.values()) / len(self.name_docs_read), 2) if self.name_docs_read else 0
                },
                "contagem_keywords_geral": {},
                "porcentagem_keywords_geral": {},
                "top_keywords_geral": [],
                "estatisticas_por_orgao": orgao_statistics,
                "documentos_lidos": self.name_docs_read,
                "imga_results": self.imga_results if self.only_integrity_plans and self.imga_results else {},
                "otimizacoes": {
                    "workers_utilizados": self.max_workers,
                    "batch_size": self.batch_size,
                    "snippets_cache_size": len(self.similarity_cache),
                    "paralelismo_ativo": True,
                    "controle_memoria_ativo": True,
                    "tempo_processamento": "calculado_em_runtime"
                }
            }

            # Processar keywords gerais
            for keyword, count in self.keywords.items():
                percentage = (count / self.total_words_read) * 100.0 if self.total_words_read > 0 else 0
                log_msg = f"{keyword} foi encontrada: {count}, que representa {percentage:.2f}%"
                self.logger.info(log_msg)
                summary["contagem_keywords_geral"][keyword] = count
                summary["porcentagem_keywords_geral"][keyword] = round(percentage, 2)

            # Top keywords gerais
            top_keywords_geral = sorted(self.keywords.items(), key=lambda x: x[1], reverse=True)[:20]
            summary["top_keywords_geral"] = [{"keyword": k, "count": c, "percentage": summary["porcentagem_keywords_geral"][k]} for k, c in top_keywords_geral]

            # Log das estatísticas por órgão
            self.logger.info("\n--- Estatísticas por Órgão ---")
            for orgao, stats in orgao_statistics.items():
                self.logger.info(f"🏢 {orgao}:")
                self.logger.info(f"   📄 Documentos: {stats['documentos_processados']}")
                self.logger.info(f"   📝 Palavras lidas: {stats['total_palavras_lidas']:,}")
                self.logger.info(f"   🔍 Total ocorrências: {stats['total_ocorrencias']}")
                self.logger.info(f"   📊 Média por documento: {stats['media_ocorrencias_por_documento']}")
                
                if stats['top_keywords']:
                    top_3 = stats['top_keywords'][:3]
                    keywords_str = ", ".join([f"{k['keyword']}({k['count']})" for k in top_3])
                    self.logger.info(f"   🏆 Top keywords: {keywords_str}")

            self.logger.info("No diretório de output você pode verificar os pedaços em que o texto foi encontrado")

            try:
                summary_path = os.path.join(os.path.dirname(self.output_path), "..", "summary_results.json")
                os.makedirs(os.path.dirname(summary_path), exist_ok=True)
                with open(summary_path, 'w', encoding='utf-8') as f:
                    json.dump(summary, f, ensure_ascii=False, indent=4)
                self.logger.info(f"Sumário detalhado dos resultados salvo em: {summary_path}")
            except Exception as e:
                self.logger.error(f"Erro ao salvar sumário: {e}")

        except Exception as e:
            self.logger.error(f"Erro ao exibir resultados: {e}")

    def get_analysed_orgaos_from_docs(self) -> List[str]:
        """Extrai os nomes dos órgãos a partir dos nomes dos documentos lidos"""
        try:
            orgaos = set()
            for doc_name in self.name_docs_read:
                try:
                    orgao = doc_name.split('_')[0]
                    orgaos.add(orgao)
                except:
                    continue
            return list(orgaos)
        except Exception as e:
            self.logger.error(f"Erro ao extrair órgãos analisados: {e}")
            return []


def company_acronym(company_name: str) -> str:
    """Retorna a sigla da empresa/órgão baseado no nome completo"""
    acronyms = {
        'Secretaria De Estado Da Agricultura, Abastecimento E Desenvolvimento Rural': 'SEAGRI',
        'Secretaria De Estado De Atendimento À Comunidade': 'SEAC',
        'Casa Civil': 'CACI',
        'Casa Militar': 'CM',
        'Secretaria De Estado De Comunicação': 'SECOM',
        'Secretaria De Estado De Cultura E Economia Criativa': 'SECEC',
        'Secretaria De Estado De Ciência, Tecnologia E Inovação': 'SECTI',
        'Secretaria De Desenvolvimento Social': 'SEDES',
        'Secretaria De Estado De Educação': 'SEEDF',
        'Secretaria De Estado De Esporte E Lazer Do Distrito Federal': 'SELDF',
        'Secretaria De Estado De Economia': 'SEEC',
        'Secretaria De Desenvolvimento Urbano E Habitação': 'SEDUH',
        'Secretaria De Estado De Justiça E Cidadania': 'SEJUS',
        'Secretaria De Estado Do Meio Ambiente E Proteção Animal': 'SEMA',
        'Secretaria De Estado Da Mulher': 'SMDF',
        'Secretaria De Estado De Obras E Infraestrutura': 'SODF',
        'Secretaria De Estado De Família E Juventude': 'SEJUV',
        'Secretaria De Estado De Projetos Especiais': 'SEPE',
        'Secretaria De Estado De Relações Institucionais': 'SERINS',
        'Secretaria De Estado De Saúde': 'SES',
        'Secretaria De Estado De Segurança Pública': 'SSPDF',
        'Secretaria De Estado De Desenvolvimento Econômico, Trabalho E Renda': 'SEDET',
        'Secretaria De Estado De Transporte E Mobilidade': 'SEMOB',
        'Secretaria De Estado De Turismo': 'SETUR',
        'Secretaria De Estado De Governo': 'SEGOV',
        'Secretaria De Estado De Proteção Da Ordem Urbanística – Df Legal': 'DF_LEGAL',
        'Secretaria De Estado De Administração Penitenciária – Seape': 'SEAPE',
        'Secretaria Da Pessoa Com Deficiência Do Distrito Federal': 'SEPD',
        'Secretaria De Estado De Assuntos Internacionais': 'SERINTER',
        'Controladoria-Geral do DF': 'CGDF',
        'Polícia Civil do DF': 'PCDF',
        'Instituto de Previdência dos Servidores do DF': 'IPREVDF',
        'Companhia do Metrapolitano do DF': 'METRO',
        'Agência Reguladora de Águas, Energia e Saneamento Básico do DF': 'ADASA',
        'Fundação de Apoio à Pesquisa do DF': 'FAPDF',
        'Companhia de Desenvolvimento Habitacional do DF': 'CODHAB',
        'Centrais de Abastecimento do DF': 'CEASA',
        'Instituto de Assistência à Saúde dos Servidores do DF': 'INAS',
        'Instituto de Pesquisa e Estatística do DF': 'IPEDF',
        'Companhia Urbanizadora da Nova Capital do Brasil': 'NOVACAP',
        'Departamento de Estradas de Rodagem do DF': 'DER',
        'Serviço de Limpeza Urbana do DF': 'SLU',
        'Instituo de Defesa do Consumidor do DF': 'PROCON',
        'Procuradoria-Geral do DF': 'PGDF',
        'Corpos de Bombeiros Militar do DF': 'CBM',
        'Arquivo Público do DF': 'ARPDF',
        'Departamento Estadual de Trânsito do DF': 'DETRAN',
        'Instituto Brasília Ambiental': 'IBRAM',
        'Companhia Imobiliária de Brasília': 'TERRACAP',
        'Banco de Brasília': 'BRB',
        'Companhia Energética de Brasília': 'CEB',
        'Fundação de Amparo ao Trabalhador Preso do DF': 'FUNAP',
        'Fundação de Ensino e Pesquisa em Ciências da Saúde': 'FEPECS',
        'Jardim Botânico de Brasília': 'JBB',
        'Fundação Jardim Zoológico de Brasília': 'FJZB',
        'Universidade do Distrito Federal': 'UNDF',
        'Junta Comercial, Industrial e Serviços do DF': 'JUCIS',
        'Polícia Militar do DF': 'PMDF',
        'Companhia de Saneamento Ambiental do Distrito Federal': 'CAESB',
        'Empresa de Assistência Técnica e Extensão Rural do DF': 'EMATER',
        'Sociedade de Transportes Coletivos de Brasília': 'TCB',
        'Fundação Hemocentro de Brasília': 'FHB',
        'Escola de Governo do Distrito Federal': 'EGOV',
    }
    return acronyms.get(company_name, "")





if __name__ == "__main__":
    try:
        batch_size = int(os.environ.get('DANI_BATCH_SIZE', 15))
        max_workers = int(os.environ.get('DANI_MAX_WORKERS', 2))

        d = Dani(batch_size=batch_size, max_workers=max_workers)
        d.run()
    except KeyboardInterrupt:
        print("\n⚠️ Processamento interrompido pelo usuário")
    except Exception as e:
        print(f"❌ Erro crítico: {e}")
        import traceback

        traceback.print_exc()
    finally:
        gc.collect()