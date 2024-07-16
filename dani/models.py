# ###################################################################
# ## DANI - Desenvolvedor e Apresentador de Números e Indicadores ##
# ##################################################################

# from os import listdir, remove
# # from os import system
# from os.path import isfile, isdir, join, exists

# from unidecode import unidecode

# from PIL import Image
# import pdf2image
# import pytesseract

# import docx

# import re

# class Dani:
#     docs_path = "dani/docs/input/"
#     is_silga = False
#     all_orgaos = ""
#     orgao_name = ""
#     name_docs_read = []
#     read_ratio = 3
#     total_words_read = 0
#     keywords = {}
#     snippets = []
#     output_path = "dani/docs/output/resultados.txt"
#     docx_path = "dani/docs/output/"

#     def __init__(self) -> None:
#         self.catch_param()

#         if (self.all_orgaos == 's'):
#             self.read_all_docs()
#         else:
#             self.read_docs()

#         self.display_results()

#         exit(0)

#     def catch_param(self) -> None:
#         self.all_orgaos = input("Você deseja fazer uma pesquisa geral (s/n): ")

#         if (self.all_orgaos == 'n'):
#             self.orgao_name = input("Insira o nome do orgão: ")#.title()
#             acronym = company_acronym(self.orgao_name.upper())
#             if not acronym == "":
#                 self.orgao_name = acronym.lower()
#             else:
#                 self.orgao_name = self.orgao_name.lower()

#             self.read_ratio = int(input("Insira a quantidade de documentos que serão lidos: "))

#         tratar_keywords = input("Insira as palavras-chaves(separadas por vírgula): ")
#         keywords = tratar_keywords.split(",")
#         for keyword in keywords:
#             self.keywords.setdefault(keyword.lower(), 0)

#     def read_docs(self) -> None:
#         self.docs_path = f"{self.docs_path}{self.orgao_name}"
#         if (self.read_ratio == 0):
#             self.read_ratio = len(listdir(self.docs_path))

#         pdfs_read = 0
#         for file in listdir(self.docs_path):
#             if isfile(join(self.docs_path, file)):
#                 # REF: colocar esta função para retornar uma lista de caminho de arquivos
#                 # E o método 'read_pdf_file' receber está lista como parâmetro
#                 self.read_pdf_file(f"{self.docs_path}/{file}", self.orgao_name);

#                 self.name_docs_read.append(file)
#                 pdfs_read += 1

#                 if pdfs_read >= self.read_ratio:
#                     break

#     def read_all_docs(self) -> None:
#         all_dir = [join(self.docs_path, d) for d in listdir(self.docs_path) if isdir(join(self.docs_path, d))]
#         self.orgao_name = [d for d in listdir(self.docs_path) if isdir(join(self.docs_path, d))]

#         file_atas_dict = {}

#         all_files = []

#         for dir in all_dir:
#             for file in listdir(dir):
#                 self.name_docs_read.append(file)
#                 all_files.append(join(dir, file))
#                 file_atas_dict.setdefault(dir.split('/')[3], []).append(join(dir, file))

#         for key, values in file_atas_dict.items():
#             for value in values:
#                 self.read_pdf_file(value, key)
        

#     def read_pdf_file(self, file_name : str, orgao_name: str) -> None:
#         images = pdf_to_image(file_name);

#         pdf_pages_text = []


#         for _, img in enumerate(images):
#             pdf_pages_text.append(ocr_core(img))

#         # reader = PyPDF2.PdfReader(file_name)

#         # number_pages = len(reader.pages)

#         # for n_page in range(number_pages):
#         #     if reader.pages[n_page].extract_text() is not None:
#         #         page_text = reader.pages[n_page].extract_text()
#         #         print(page_text)
#         #         pdf_pages_text.append(page_text)

#         for page_text in pdf_pages_text:
#             self.total_words_read += len(page_text.split())

#         snippets = []
#         doc = self.create_docx(orgao_name)

#         for page_text in pdf_pages_text:
#             keyword_in_page = self.count_keywords(page_text)

#             if keyword_in_page:
#                 for keyword in self.keywords:
#                     if self.keywords[keyword] != 0:
#                         snippets = self.save_text_snippet(page_text, keyword)
#                         if snippets is not None:
#                             # self.save_snippets_file(snippets)
#                             for snippet in snippets:
#                                 self.write_docx(doc, snippet, keyword)
    
#         new_docx_path = join(self.docx_path, 'output' + orgao_name + '.docx')
#         doc.save(new_docx_path)

#     def display_results(self) -> None:
#         print("Os orgãos que tiveram suas atas do CIG lidos foram:")
#         print(self.orgao_name)
        
#         print("Os documentos lidos foram:")
#         print(self.name_docs_read)
            
#         print("Quantidade de vezes que cada palavra chave foi encontrada:")
#         for keyword in self.keywords:
#             keyword_count = self.keywords[keyword]
#             print(f"{keyword} foi encontrada: {keyword_count}, que representa {(keyword_count / self.total_words_read) * 100.0}")

#         print("No diretório \output você pode verificar o pedaços em que o texto foi encontrado")    

#     def save_text_snippet(self, page_text: str, keyword: str, window_size: int = 210) -> list:
#         snipppets = []
#         index = 0

#         while index < len(page_text) :
#             found_index = page_text.find(keyword.lower(), index)

#             if found_index == -1:
#                 break

#             start_index = max(0, found_index - window_size)
#             end_index = min(len(page_text), found_index + len(keyword) + window_size)

#             snipppet = page_text[start_index:end_index]
#             snipppet = '"' + snipppet + '"'
#             snipppets.append(snipppet)
#             index = found_index + len(keyword)

#         return snipppets

#     # def save_snippets_file(self, snippets: list) -> None:
#     #     with open(self.output_path, 'w') as file:
#     #         for snippet in snippets:
#     #             file.write(snippet + '\n\n')

#     def create_docx(self, orgao_name: str) -> docx.Document:
#         new_docx_path = join(self.docx_path, 'output' + orgao_name + '.docx')
#         if isfile(self.docx_path):
#             doc = docx.Document(docx=new_docx_path)
#         else:
#             doc = docx.Document()

#             doc.add_heading('Resultado Dani', 0)
#             doc.add_heading(f'Análise atas CIG {orgao_name}', 2)

#         return doc

#     def write_docx(self, doc: docx.Document, snippet: str, keyword: str) -> None:
#         doc.add_heading(f'Palavra-chave: {keyword}', 3)

#         p = doc.add_paragraph()
#         p.paragraph_format.line_spacing = 1
#         p.paragraph_format.space_after = 0

#         valid_chars = ''.join(c for c in snippet if c.isprintable())
#         snippet_fixed = ''.join(c if c.isprintable() else ' ' for c in snippet)

#         p.add_run(snippet_fixed)
#         p.add_run("\n")

#     def count_keywords(self, text: str):    
#         verif = False
#         for keyword in self.keywords:
#             default = re.compile(r'\b' + re.escape(keyword) + r'\w*\b', re.IGNORECASE)
#             self.keywords[keyword] += len(default.findall(text))

#             if len(default.findall(text)) != 0:
#                 verif = True

#         return verif

# def company_acronym(company_name: str) -> str:
#     acronyms = {
#         'Secretaria De Estado Da Agricultura, Abastecimento E Desenvolvimento Rural': 'SEAGRI',
#         'Secretaria De Estado De Atendimento À Comunidade': 'SEAC',
#         'Casa Civil': 'CACI',
#         'Casa Militar': 'CM',
#         'Secretaria De Estado De Comunicação': 'SECOM',
#         'Secretaria De Estado De Cultura E Economia Criativa': 'SECEC',
#         'Secretaria De Estado De Ciência, Tecnologia E Inovação': 'SECTI',
#         'Secretaria De Desenvolvimento Social': 'SEDES',
#         'Secretaria De Estado De Educação': 'SEE',
#         'Secretaria De Estado De Esporte E Lazer Do Distrito Federal': 'SELDF',
#         'Secretaria De Estado De Economia': 'SEEC',
#         'Secretaria De Desenvolvimento Urbano E Habitação': 'SEDUH',
#         'Secretaria De Estado De Justiça E Cidadania': 'SEJUS',
#         'Secretaria De Estado Do Meio Ambiente E Proteção Animal': 'SEMA',
#         'Secretaria De Estado Da Mulher': 'SMDF',
#         'Secretaria De Estado De Obras E Infraestrutura': 'SODF',
#         'Secretaria De Estado De Família E Juventude': 'SEJUV',
#         'Secretaria De Estado De Projetos Especiais': 'SEPE',
#         'Secretaria De Estado De Relações Institucionais': 'SERINS',
#         'Secretaria De Estado De Saúde': 'SES',
#         'Secretaria De Estado De Segurança Pública': 'SSP/DF',
#         'Secretaria De Estado De Desenvolvimento Econômico, Trabalho E Renda': 'SEDET',
#         'Secretaria De Estado De Transporte E Mobilidade': 'SEMOB',
#         'Secretaria De Estado De Turismo': 'SETUR',
#         'Secretaria De Estado De Governo': 'SEGOV',
#         'Secretaria De Estado De Proteção Da Ordem Urbanística – Df Legal': 'DF LEGAL',
#         'Secretaria De Estado De Administração Penitenciária – Seape': 'SEAPE',
#         'Secretaria Da Pessoa Com Deficiência Do Distrito Federal': 'SEPD',
#         'Secretaria De Estado De Assuntos Internacionais': 'SERINTER',
#         'Controladoria-Geral do DF': 'CGDF',
#         'Procuradoria Geral do DF': 'PGDF',
#         'Polícia Civil': 'PCDF',
#         'Polícia Militar': 'PMDF',
#         'Corpo de Bombeiros': 'CBMDF',
#         'Arquivo Público Do Distrito Federal': 'ARPDF',
#         'Agência Reguladora de Águas e Saneamento do DF': 'ADASA',
#         'Departamento de Estradas de Rodagem do DF': 'DER',
#         'Departamento de Trânsito do DF': 'DETRAN',
#         'Instituto de Assistência à Saúde dos Servidores do DF': 'INAS',
#         'Instituto de Defesa do Consumidor do DF': 'PROCON-DF',
#         'Instituto de Previdência dos Servidores do DF': 'IPREV',
#         'Instituto Brasília Ambiental': 'IBRAM',
#         'Serviço de Limpeza Urbana do DF': 'SLU',
#         'Centrais de Abastecimento do Distrito Federal': 'CEASA',
#         'Companhia de Desenvolvimento Habitacional': 'CODHAB-DF',
#         'Companhia Imobiliária de Brasília': 'TERRACAP',
#         'Instituto de Pesquisa e Estatística do Distrito Federal': 'IPEDF',
#         'Companhia de Saneamento Ambiental do DF': 'CAESB',
#         'Companhia Urbanizadora da Nova Capital do Brasil': 'NOVACAP',
#         'Empresa de Assistência Técnica e Extensão Rural': 'EMATER',
#         'Companhia do Metropolitano do Distrito Federal': 'METRO-DF',
#         'Sociedade de Transportes Coletivos de Brasília': 'TCB',
#         'Banco de Brasília': 'BRB',
#         'Companhia Energética de Brasília': 'CEB',
#         'Fundação de Amparo ao Trabalhador Preso do Distrito Federal': 'FUNAP',
#         'Fundação de Apoio à Pesquisa do Distrito Federal': 'FAPDF',
#         'Fundação de Ensino e Pesquisa em Ciências da Saúde': 'FEPECS',
#         'Fundação Hemocentro de Brasília': 'FHB',
#         'Fundação Jardim Zoológico de Brasília': 'FJZB',
#         'Fundação Jardim Botânico': 'JBB',
#         'Universidade do Distrito Federal': 'UNDF',
#         'Escola de Governo': 'EGOV',
#         'Junta Comercial, Industrial E Serviços Do Distrito Federal': 'JUCIS-DF',
#         'Secretaria De Estado De Planejamento, Orçamento E Administração Do Distrito Federal': 'SEPLAD',
#     }


#     acronym = company_name in acronyms
#     if acronym:
#         return acronyms[company_name]
#     else:
#         return ""
    
# def pdf_to_image(pdf_file):
#     return pdf2image.convert_from_path(pdf_file)
    

# def ocr_core(file):
#     text = pytesseract.image_to_string(file, lang='por')
#     return text.lower()

###################################################################
## DANI - Desenvolvedor e Apresentador de Números e Indicadores ##
##################################################################

from os import listdir, remove
# from os import system
from os.path import isfile, isdir, join, exists

from unidecode import unidecode

from PIL import Image
import pdf2image
import pytesseract

import docx

import re

class Dani:
    docs_path = "dani/docs/input/"
    is_silga = False
    all_orgaos = ""
    orgao_name = ""
    name_docs_read = []
    read_ratio = 3
    total_words_read = 0
    keywords = {}
    snippets = []
    output_path = "dani/docs/output/resultados.txt"
    docx_path = "dani/docs/output/"

    def __init__(self) -> None:
        self.docx_files = {}  # Para armazenar documentos abertos por palavra-chave
        self.catch_param()

        if (self.all_orgaos == 's'):
            self.read_all_docs()
        else:
            self.read_docs()

        self.save_and_close_docs()
        self.display_results()

        exit(0)

    def catch_param(self) -> None:
        self.all_orgaos = input("Você deseja fazer uma pesquisa geral (s/n): ")

        if (self.all_orgaos == 'n'):
            self.orgao_name = input("Insira o nome do orgão: ")#.title()
            acronym = company_acronym(self.orgao_name.upper())
            if not acronym == "":
                self.orgao_name = acronym.lower()
            else:
                self.orgao_name = self.orgao_name.lower()

            self.read_ratio = int(input("Insira a quantidade de documentos que serão lidos: "))

        tratar_keywords = input("Insira as palavras-chaves(separadas por vírgula): ")
        keywords = tratar_keywords.split(",")
        for keyword in keywords:
            self.keywords.setdefault(keyword.lower(), 0)

    def read_docs(self) -> None:
        self.docs_path = f"{self.docs_path}{self.orgao_name}"
        if (self.read_ratio == 0):
            self.read_ratio = len(listdir(self.docs_path))

        pdfs_read = 0
        for file in listdir(self.docs_path):
            if isfile(join(self.docs_path, file)):
                self.read_pdf_file(f"{self.docs_path}/{file}", self.orgao_name)

                self.name_docs_read.append(file)
                pdfs_read += 1

                if pdfs_read >= self.read_ratio:
                    break

    def read_all_docs(self) -> None:
        all_dir = [join(self.docs_path, d) for d in listdir(self.docs_path) if isdir(join(self.docs_path, d))]
        self.orgao_name = [d for d in listdir(self.docs_path) if isdir(join(self.docs_path, d))]

        file_atas_dict = {}

        all_files = []

        for dir in all_dir:
            for file in listdir(dir):
                self.name_docs_read.append(file)
                all_files.append(join(dir, file))
                file_atas_dict.setdefault(dir.split('/')[3], []).append(join(dir, file))

        for key, values in file_atas_dict.items():
            for value in values:
                self.read_pdf_file(value, key)

    def read_pdf_file(self, file_name: str, orgao_name: str) -> None:
        images = pdf_to_image(file_name)
        pdf_pages_text = []

        for _, img in enumerate(images):
            pdf_pages_text.append(ocr_core(img))

        for page_text in pdf_pages_text:
            self.total_words_read += len(page_text.split())

        snippets = []

        for page_text in pdf_pages_text:
            keyword_in_page = self.count_keywords(page_text)

            if keyword_in_page:
                for keyword in self.keywords:
                    if self.keywords[keyword] != 0:
                        snippets = self.save_text_snippet(page_text, keyword)
                        if snippets is not None:
                            for snippet in snippets:
                                self.write_keyword_docx(snippet, keyword, orgao_name)

    def display_results(self) -> None:
        print("Os orgãos que tiveram suas atas do CIG lidos foram:")
        print(self.orgao_name)
        
        print("Os documentos lidos foram:")
        print(self.name_docs_read)
            
        print("Quantidade de vezes que cada palavra chave foi encontrada:")
        for keyword in self.keywords:
            keyword_count = self.keywords[keyword]
            print(f"{keyword} foi encontrada: {keyword_count}, que representa {(keyword_count / self.total_words_read) * 100.0}")

        print("No diretório \output você pode verificar o pedaços em que o texto foi encontrado")    

    def save_text_snippet(self, page_text: str, keyword: str, window_size: int = 210) -> list:
        snipppets = []
        index = 0

        while index < len(page_text):
            found_index = page_text.find(keyword.lower(), index)

            if found_index == -1:
                break

            start_index = max(0, found_index - window_size)
            end_index = min(len(page_text), found_index + len(keyword) + window_size)

            snipppet = page_text[start_index:end_index]
            snipppet = '"' + snipppet + '"'
            snipppets.append(snipppet)
            index = found_index + len(keyword)

        return snipppets

    def create_keyword_docx(self, keyword: str, orgao_name: str) -> docx.Document:
        new_docx_path = join(self.docx_path, f'output_{orgao_name}_{keyword}.docx')
        if isfile(new_docx_path):
            doc = docx.Document(new_docx_path)
        else:
            doc = docx.Document()
            doc.add_heading('Resultado Dani', 0)
            doc.add_heading(f'Análise atas CIG {orgao_name}', 2)
            doc.add_heading(f'Palavra-chave: {keyword}', 2)
        return doc

    def write_keyword_docx(self, snippet: str, keyword: str, orgao_name: str) -> None:
        if keyword not in self.docx_files:
            self.docx_files[keyword] = self.create_keyword_docx(keyword, orgao_name)

        doc = self.docx_files[keyword]

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = 0

        valid_chars = ''.join(c for c in snippet if c.isprintable())
        snippet_fixed = ''.join(c if c.isprintable() else ' ' for c in snippet)

        p.add_run(snippet_fixed)
        p.add_run("\n")

    def save_and_close_docs(self) -> None:
        for keyword, doc in self.docx_files.items():
            new_docx_path = join(self.docx_path, f'output_{self.orgao_name}_{keyword}.docx')
            doc.save(new_docx_path)

    def count_keywords(self, text: str) -> bool:
        verif = False
        for keyword in self.keywords:
            default = re.compile(r'\b' + re.escape(keyword) + r'\w*\b', re.IGNORECASE)
            self.keywords[keyword] += len(default.findall(text))

            if len(default.findall(text)) != 0:
                verif = True

        return verif

def company_acronym(company_name: str) -> str:
    acronyms = {
        'Secretaria De Estado Da Agricultura, Abastecimento E Desenvolvimento Rural': 'SEAGRI',
        'Secretaria De Estado De Atendimento À Comunidade': 'SEAC',
        'Casa Civil': 'CACI',
        'Casa Militar': 'CM',
        'Secretaria De Estado De Comunicação': 'SECOM',
        'Secretaria De Estado De Cultura E Economia Criativa': 'SECEC',
        'Secretaria De Estado De Ciência, Tecnologia E Inovação': 'SECTI',
        'Secretaria De Desenvolvimento Social': 'SEDES',
        'Secretaria De Estado De Educação': 'SEE',
        'Secretaria De Estado De Esporte E Lazer Do Distrito Federal': 'SELDF',
        'Secretaria De Estado De Economia': 'SEEC',
        'Secretaria De Desenvolvimento Urbano E Habitação': 'SEDUH',
        'Secretaria De Estado De Justiça E Cidadania': 'SEJUS',
        'Secretaria De Estado Do Meio Ambiente E Proteção Animal': 'SEMA',
        'Secretaria De Estado Da Mulher': 'SMDF',
        'Secretaria De Estado De Obras E Infraestrutura': 'SODF',
        'Secretaria De Estado De Família E Juventude': 'SEJUV',
        'Secretaria De Estado De Projetos Especiais': 'SEPE',
        'Secretaria De Estado De Relações Institucionais': 'SERINS',
        'Secretaria De Estado De Saúde': 'SES',
        'Secretaria De Estado De Segurança Pública': 'SSP/DF',
        'Secretaria De Estado De Desenvolvimento Econômico, Trabalho E Renda': 'SEDET',
        'Secretaria De Estado De Transporte E Mobilidade': 'SEMOB',
        'Secretaria De Estado De Turismo': 'SETUR',
        'Secretaria De Estado De Governo': 'SEGOV',
        'Secretaria De Estado De Proteção Da Ordem Urbanística – Df Legal': 'DF LEGAL',
        'Secretaria De Estado De Administração Penitenciária – Seape': 'SEAPE',
        'Secretaria Da Pessoa Com Deficiência Do Distrito Federal': 'SEPD',
        'Secretaria De Estado De Assuntos Internacionais': 'SERINTER',
        'Controladoria-Geral do DF': 'CGDF'
    }

    acronym = company_name in acronyms
    if acronym:
        return acronyms[company_name]
    else:
        return ""
    
def pdf_to_image(pdf_file):
    return pdf2image.convert_from_path(pdf_file)
    

def ocr_core(file):
    text = pytesseract.image_to_string(file, lang='por')
    return text.lower()